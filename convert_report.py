import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def convert_md_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # Define some basic styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue

        # Headings
        if line.startswith('# '):
            h = doc.add_heading(line[2:], level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        
        # Tables
        elif line.startswith('|'):
            # Basic table support
            table_data = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                t_line = lines[i].strip()
                if not re.match(r'^|[ \-:]+\|', t_line): # Skip separator row
                    cells = [c.strip() for c in t_line.split('|') if c.strip()]
                    if cells:
                        table_data.append(cells)
                i += 1
            
            if table_data:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                for row_idx, row_cells in enumerate(table_data):
                    for col_idx, cell_text in enumerate(row_cells):
                        table.cell(row_idx, col_idx).text = cell_text
            continue

        # Lists
        elif line.startswith('* ') or line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^- \[[x ]\]', line): # Task list
            status = " [TAMAM]" if '[x]' in line else " [BEKLİYOR]"
            content = line[6:].strip()
            doc.add_paragraph(content + status, style='List Bullet')
        
        # Normal text
        else:
            # Simple bold/italic handling
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True
                elif part.startswith('*') and part.endswith('*'):
                    p.add_run(part[1:-1]).italic = True
                else:
                    p.add_run(part)

        i += 1

    doc.save(docx_path)
    print(f"Successfully converted {md_path} to {docx_path}")

if __name__ == "__main__":
    MD_FILE = r'C:\Users\Ahmet\.gemini\antigravity\brain\817d54ce-3607-4f8b-9e9f-491776d8ea8f\squadup_genel_rapor.md'
    DOCX_FILE = r'C:\Users\Ahmet\.gemini\antigravity\scratch\squadup\SquadUp_Genel_Raporu_v2.docx'
    convert_md_to_docx(MD_FILE, DOCX_FILE)
