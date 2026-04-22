# -*- coding: utf-8 -*-
"""
SquadUp - Teknoloji Secimi ve Gerceklendirme Dokumani
5. Hafta Fonksiyonel Programlama Lab Foyu - DOCX Uretici
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"c:\Users\Ahmet\.gemini\antigravity\scratch\squadup\mobile\SquadUp_Teknoloji_Secimi_Raporu.docx"

BLUE_RGB    = RGBColor(0x1a, 0x36, 0x5d)
LBLUE_RGB   = RGBColor(0x2b, 0x6c, 0xb0)
GRAY_RGB    = RGBColor(0x71, 0x80, 0x96)
GREEN_RGB   = RGBColor(0x27, 0x67, 0x49)
WHITE_RGB   = RGBColor(0xFF, 0xFF, 0xFF)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{{edge}}".replace("{edge}", edge))
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:color"), "C0C0C0")
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def add_heading(doc, text, level=1, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level==1 else 8)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14 if level==1 else 11)
    run.font.color.rgb = color or (BLUE_RGB if level==1 else LBLUE_RGB)
    if level == 1:
        # underline via border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:color"), "1a365d")
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p

def add_body(doc, text, justify=True):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(2)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0x4a, 0x55, 0x68)
    return p

def add_bullet(doc, text, check=False):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)
    sym = "✓" if check else "•"
    run = p.add_run(f"{sym}  {text}")
    run.font.size = Pt(10)
    run.font.color.rgb = GREEN_RGB if check else RGBColor(0x4a, 0x55, 0x68)
    return p

def make_table(doc, headers, rows, col_widths):
    total_cols = len(headers)
    table = doc.add_table(rows=1+len(rows), cols=total_cols)
    table.style = "Table Grid"

    # Header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.width = col_widths[i]
        set_cell_bg(cell, "1a365d")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = WHITE_RGB

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri+1]
        bg = "ffffff" if ri % 2 == 0 else "f7f7f7"
        for ci, text in enumerate(row_data):
            cell = row.cells[ci]
            cell.width = col_widths[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(text)
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x4a, 0x55, 0x68)
            if ci == 0:
                run.bold = True
                run.font.color.rgb = LBLUE_RGB

    return table

def build_docx():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)
        section.top_margin    = Cm(2.2)
        section.bottom_margin = Cm(2.2)

    # ── HEADER BOX ───────────────────────────────────────────────────────────
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, "1a365d")
    cell.width = Cm(16.6)

    for text, sz, color in [
        ("SquadUp", 22, "FFFFFF"),
        ("Teknoloji Secimi ve Gerceklendirme Dokumani", 13, "BEE3F8"),
        ("5. Hafta Fonksiyonel Programlama Laboratuvari", 10, "BEE3F8"),
        ("[Ogrenci Adi — Ogrenci Numarasi]  •  2025", 9, "A0AEC0"),
    ]:
        p = cell.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        run = p.runs[0]
        run.font.size = Pt(sz)
        run.bold = (sz >= 13)
        r, g, b = int(color[0:2],16), int(color[2:4],16), int(color[4:6],16)
        run.font.color.rgb = RGBColor(r, g, b)

    doc.add_paragraph()

    # ── GIRIS ────────────────────────────────────────────────────────────────
    add_heading(doc, "Giris")
    add_body(doc,
        "Bu dokuman, SquadUp mobil uygulamasinin gelistirilmesinde kullanilacak "
        "teknolojilerin secimini ve bu secimlerin gerceklendirmesini icermektedir. "
        "Secimler; problemin yapisi, kullanici sayisi beklentisi ve gelistirme suresi "
        "dikkate alinarak yapilmistir. \"Bildigim teknolojiyi kullandim\" yaklasiminin "
        "yerine \"probleme uygun teknolojiyi sectim\" yaklasimi benimsenmistir.")

    doc.add_paragraph()

    # ── OZET TABLO ───────────────────────────────────────────────────────────
    add_heading(doc, "Teknoloji Secim Ozeti")
    make_table(doc,
        headers=["Baslik", "Secim"],
        rows=[
            ["2.1  Mobil Gelistirme", "Cross-platform — React Native"],
            ["2.2  Veri Kaynagi",     "Bulut tabanli (Flask REST API + SQLite)"],
            ["2.3  Ek Kutuphaneler",  "Zustand, React Navigation, Expo, TypeScript"],
        ],
        col_widths=[Cm(5.5), Cm(11.1)]
    )

    doc.add_paragraph()

    # ── 2.1 ──────────────────────────────────────────────────────────────────
    add_heading(doc, "2.1  Mobil Gelistirme Teknolojisi: React Native (Cross-platform)", level=2)
    add_body(doc,
        "SquadUp, hem Android hem iOS kullanicilarina ulasmayi hedefleyen bir aktivite "
        "eslestirme platformudur. Bu nedenle cross-platform yaklasimi secilmistir. "
        "React Native; tek bir kod tabaniyla her iki platforma native performans sunmakta, "
        "JavaScript/TypeScript bilgisiyle hizla gelistirilebilmekte ve genis bir ekosisteme "
        "sahip bulunmaktadir.")

    make_table(doc,
        headers=["Gerceklendirme Kriteri", "Aciklama"],
        rows=[
            ["Problem yapisi",    "Kullanici eslestirme mantigi platform bagimsizdir"],
            ["Kullanici sayisi",  "Android + iOS pazar paylasimi icin tek kod tabani yeterlidir"],
            ["Gelistirme suresi", "Ayri native uygulamaya gore yaklasik 2x daha hizlidir"],
        ],
        col_widths=[Cm(5.5), Cm(11.1)]
    )

    doc.add_paragraph()

    # ── 2.2 ──────────────────────────────────────────────────────────────────
    add_heading(doc, "2.2  Veri Kaynagi: Bulut Tabanli Veritabani (Flask REST API)", level=2)
    add_body(doc,
        "Uygulama, birden fazla kullanicinin ayni veri setini paylastigi bir sosyal platform "
        "oldugu icin bulut tabanli veri kaynagi zorunludur. Backend olarak Flask REST API, "
        "veri deposu olarak SQLite (gelistirme) ve ilerleyen asamada PostgreSQL kullanilmaktadir. "
        "Mobil uygulama bu API'ye fetch() / async-await ile baglanmaktadir.")

    for txt in [
        "Birden fazla kullanicinin gercek zamanli eslesmesi icin merkezi veri zorunludur",
        "SQLite - PostgreSQL gecisi uygulama katmanini etkilemez (ORM soyutlamasi)",
        "Offline-first: yerel Zustand state'i ile anlik filtreleme desteklenir",
    ]:
        add_bullet(doc, txt, check=True)

    doc.add_paragraph()

    # ── 2.3 ──────────────────────────────────────────────────────────────────
    add_heading(doc, "2.3  Ek Araclar ve Kutuphaneler", level=2)
    add_body(doc, "Asagidaki ek kutuphaneler kullanilmaktadir:")

    make_table(doc,
        headers=["Kutuphane", "Kullanim Amaci"],
        rows=[
            ["Zustand",           "Immutable state yonetimi — Redux'a gore minimal ve performansli"],
            ["React Navigation",  "Ekranlar arasi gecis (Bottom Tab Navigator + Stack)"],
            ["Expo",              "Hizli prototip ve cross-platform build ortami"],
            ["TypeScript",        "Tip guvenligi — derleme zamaninda hata yakalama, refactoring destegi"],
        ],
        col_widths=[Cm(4), Cm(12.6)]
    )

    doc.add_paragraph()

    # ── 2.4 SONUC ─────────────────────────────────────────────────────────────
    add_heading(doc, "2.4  Genel Gerceklendirme ve Sonuc")
    add_body(doc,
        "Secilen teknoloji seti, SquadUp'in temel gereksinimlerini karsilamak uzere "
        "ozenle belirlenmistir. Her secim, problemin somut ihtiyaclarindan kaynaklanmaktadir:")

    for txt in [
        "Cross-platform React Native: Sinirli gelistirme suresinde maksimum platform kapsami",
        "Bulut API: Cok kullanicili sosyal platform icin veri tutarliligi zorunlulugu",
        "Zustand: Kucuk, okunabilir ve test edilebilir state yonetimi",
        "TypeScript: Buyuyen kod tabaninda guvenli refactoring ve tip denetimi",
    ]:
        add_bullet(doc, txt, check=True)

    # Footer line
    doc.add_paragraph()
    p = doc.add_paragraph("Teslim formati: DOCX / PDF  •  Onerilen uzunluk: 1-2 sayfa  •  2025")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = GRAY_RGB

    doc.save(OUTPUT)
    print("DOCX olusturuldu:", OUTPUT)

if __name__ == "__main__":
    build_docx()
