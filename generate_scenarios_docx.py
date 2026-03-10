from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_scenarios_docx():
    doc = Document()

    # Title
    title = doc.add_heading('SquadUp – Kullanıcı Senaryoları (LAB 3)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run('Proje Adı: ').bold = True
    meta.add_run('SquadUp – Akıllı Aktivite Eşleştirme Sistemi\n')
    meta.add_run('Tarih: ').bold = True
    meta.add_run('Mart 2026')

    doc.add_paragraph()

    # Scenario 1
    doc.add_heading('Senaryo 1: Profil Oluşturma', level=1)
    p1 = doc.add_paragraph()
    p1.add_run('Kullanıcı: ').bold = True
    p1.add_run('Yeni üniversite öğrencisi\n')
    p1.add_run('Amaç: ').bold = True
    p1.add_run('İlgi alanlarını ve yetenek seviyelerini sisteme girmek.')

    steps1 = [
        "Kullanıcı uygulamayı açar ve kayıt olur.",
        "Karşılama ekranında 'Profilimi Düzenle' butonuna dokunur.",
        "İlgi duyduğu spor dallarını (örneğin basketbol, tenis) seçer.",
        "Her spor dalı için kendi seviyesini (başlangıç, orta, ileri) belirler.",
        "'Kaydet' butonuna basarak bilgilerini onaylar.",
        "Sistem profili günceller ve ana sayfaya yönlendirir.",
        "Kullanıcı, seçtiği alanlara uygun ilk önerileri ana sayfada görür."
    ]
    for i, step in enumerate(steps1, 1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_paragraph()

    # Scenario 2
    doc.add_heading('Senaryo 2: Etkinlik Düzenleme', level=1)
    p2 = doc.add_paragraph()
    p2.add_run('Kullanıcı: ').bold = True
    p2.add_run('Sosyal bir öğrenci\n')
    p2.add_run('Amaç: ').bold = True
    p2.add_run('Hafta sonu için bir basketbol maçı düzenlemek ve oyuncu toplamak.')

    steps2 = [
        "Kullanıcı ana ekranda 'Etkinlik Oluştur' seçeneğini seçer.",
        "Etkinlik türü olarak 'Basketbol'u işaretler.",
        "Maçın tarihini, saatini ve yerini girer.",
        "Aranan oyuncuların beklenen seviyesini 'Orta Seviye' olarak belirler.",
        "Etkinlik için kısa bir açıklama yazar.",
        "'Yayınla' butonuna dokunarak etkinliği oluşturur.",
        "Etkinlik, diğer kullanıcıların listesinde ve önerilerinde görünür."
    ]
    for i, step in enumerate(steps2, 1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_paragraph()

    # Scenario 3
    doc.add_heading('Senaryo 3: Etkinliğe Katılma', level=1)
    p3 = doc.add_paragraph()
    p3.add_run('Kullanıcı: ').bold = True
    p3.add_run('Spor yapmak isteyen öğrenci\n')
    p3.add_run('Amaç: ').bold = True
    p3.add_run('Kendi seviyesine uygun mevcut bir etkinliğe dahil olmak.')

    steps3 = [
        "Kullanıcı uygulamayı açar ve 'Önerilen Etkinlikler' sekmesine gider.",
        "Kendi seviyesiyle %90 eşleşen bir tenis antrenmanını görür.",
        "Etkinlik detaylarına tıklayarak saati ve katılımcıları kontrol eder.",
        "'Katıl' butonuna dokunur.",
        "Sistem kullanıcıyı katılımcı listesine ekler.",
        "Kullanıcıya katılımının onaylandığına dair bir bildirim gösterilir.",
        "Etkinlik, kullanıcının 'Takvimim' bölümünde listelenir."
    ]
    for i, step in enumerate(steps3, 1):
        doc.add_paragraph(f"{i}. {step}")

    output_path = r'c:\Users\Ahmet\.gemini\antigravity\scratch\squadup\Kullanıcı_Senaryoları.docx'
    doc.save(output_path)
    print(f'Saved to: {output_path}')

if __name__ == "__main__":
    create_scenarios_docx()
